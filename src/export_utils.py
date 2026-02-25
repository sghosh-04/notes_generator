from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document


# -----------------------------
# PDF Export
# -----------------------------
def export_to_pdf(filename, content):
    """
    Exports text content to a PDF file.
    """

    styles = getSampleStyleSheet()
    style = styles['BodyText']

    doc = SimpleDocTemplate(filename, pagesize=A4)
    elements = []

    for line in content.split("\n"):
        elements.append(Paragraph(line, style))
        elements.append(Spacer(1, 6))

    doc.build(elements)


# -----------------------------
# DOCX Export
# -----------------------------
def export_to_docx(filename, content):
    """
    Exports text content to a DOCX file.
    """

    doc = Document()
    doc.add_heading("Generated Notes", level=1)

    for line in content.split("\n"):
        doc.add_paragraph(line)

    doc.save(filename)